"""
Document Q&A Web Application
Multi-format, multimodal RAG with cross-encoder re-ranking.
Supports PDF, DOCX, TXT, Markdown. Images from PDF + DOCX.
"""

import os, re, io, json, base64, hashlib, time, tempfile
from dataclasses import dataclass
from typing import List, Optional, Tuple, Any
from pathlib import Path

from flask import Flask, render_template, request, jsonify
import numpy as np, faiss
from sentence_transformers import SentenceTransformer, CrossEncoder
from rank_bm25 import BM25Okapi
from openai import OpenAI
import tiktoken

try:
    import fitz
except ImportError:
    fitz = None
try:
    import docx
except ImportError:
    docx = None


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
def get_file_type(fn):
    ext = Path(fn).suffix.lower()
    return ext if ext in SUPPORTED_EXTENSIONS else None


# ══════════════════════════════════════════════════════════════
# Data Structures
# ══════════════════════════════════════════════════════════════

@dataclass
class PageElement:
    text: str; page_num: int; element_type: str
    section_title: str = ""; image_b64: str = ""

@dataclass
class Chunk:
    chunk_id: str; text: str; page_numbers: List[int]; section_title: str
    token_count: int; chunk_index: int; has_image: bool = False
    def citation(self):
        pages = sorted(set(self.page_numbers))
        p = f"p. {pages[0]}" if len(pages) == 1 else f"pp. {pages[0]}-{pages[-1]}"
        sec = f' — "{self.section_title}"' if self.section_title else ""
        img = " [image]" if self.has_image else ""
        return f"[{p}{sec}{img}]"


# ══════════════════════════════════════════════════════════════
# Image Captioning (PDF + DOCX)
# ══════════════════════════════════════════════════════════════

IMAGE_CAPTION_PROMPT = (
    "Describe this image from a document in detail. "
    "If it is a chart/graph/table, describe data, trends, labels, axes, key values. "
    "Be specific about numbers, percentages, or text visible.")

MIME_MAP = {"png":"image/png","jpeg":"image/jpeg","jpg":"image/jpeg",
            "gif":"image/gif","bmp":"image/bmp","tiff":"image/tiff","webp":"image/webp"}

def caption_image(image_bytes, api_key, page_num, model="gpt-4o-mini", mime="image/png"):
    if not api_key or len(image_bytes) < 5000: return None
    b64 = base64.b64encode(image_bytes).decode("utf-8")
    try:
        resp = OpenAI(api_key=api_key).chat.completions.create(
            model=model, temperature=0.0, max_tokens=300,
            messages=[{"role":"user","content":[
                {"type":"text","text":IMAGE_CAPTION_PROMPT},
                {"type":"image_url","image_url":{"url":f"data:{mime};base64,{b64}","detail":"high"}}]}])
        return f"[Image on page {page_num}]: {resp.choices[0].message.content}"
    except Exception as e:
        print(f"  ⚠️ Caption failed p.{page_num}: {e}"); return None


# ══════════════════════════════════════════════════════════════
# PDF Ingestion
# ══════════════════════════════════════════════════════════════

def _pdf_images(doc, page, pn, api_key, min_w=100, min_h=100):
    elems = []
    for info in page.get_images(full=True):
        try:
            bi = doc.extract_image(info[0])
            if not bi or bi.get("width",0)<min_w or bi.get("height",0)<min_h: continue
            ext = bi.get("ext","png"); mime = MIME_MAP.get(ext,"image/png")
            cap = caption_image(bi["image"], api_key, pn, mime=mime)
            if cap:
                b64 = base64.b64encode(bi["image"]).decode("utf-8")
                elems.append(PageElement(cap, pn, "image", image_b64=f"data:{mime};base64,{b64}"))
        except: continue
    return elems

def _median_fs(page):
    sizes = []
    for b in page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]:
        if b.get("type",1)!=0: continue
        for l in b.get("lines",[]):
            for s in l.get("spans",[]):
                t = s.get("text","").strip()
                if len(t)>2: sizes.extend([s["size"]]*len(t))
    return float(np.median(sizes)) if sizes else 12.0

def _is_heading_pdf(block, med):
    lines = block.get("lines",[])
    if not lines: return False
    fs, bold = [], False
    for l in lines:
        for s in l.get("spans",[]):
            fs.append(s["size"])
            if "bold" in s.get("font","").lower(): bold=True
    if not fs: return False
    r = max(fs)/med if med>0 else 1.0
    txt = " ".join(s["text"] for l in lines for s in l.get("spans",[])).strip()
    if len(txt)>200 or len(txt)<2: return False
    return r>=1.5 or (r>=1.3 and bold) or r>=1.2 or (bold and len(txt)<80)

def ingest_pdf(fp, api_key="", extract_images=True):
    doc = fitz.open(fp); elems, sec, ic = [], "", 0
    meta = {"title":doc.metadata.get("title",""),"total_pages":len(doc),
            "file_name":os.path.basename(fp),"file_type":"pdf"}
    for pi in range(len(doc)):
        page = doc[pi]; pn = pi+1; med = _median_fs(page)
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        try:
            for t in page.find_tables():
                df = t.to_pandas()
                if not df.empty: elems.append(PageElement(df.to_markdown(index=False),pn,"table",sec))
        except: pass
        for b in blocks:
            if b.get("type",1)!=0: continue
            txt = ""
            for l in b.get("lines",[]): txt+="".join(s["text"] for s in l.get("spans",[]))+"\n"
            txt = txt.strip()
            if not txt or len(txt)<3: continue
            if _is_heading_pdf(b, med): sec=txt.strip(); elems.append(PageElement(txt,pn,"heading",sec))
            else: elems.append(PageElement(txt,pn,"paragraph",sec))
        if extract_images and api_key:
            for ie in _pdf_images(doc,page,pn,api_key):
                ie.section_title=sec; elems.append(ie); ic+=1
    doc.close(); meta["images_extracted"]=ic; return elems, meta


# ══════════════════════════════════════════════════════════════
# DOCX Ingestion
# ══════════════════════════════════════════════════════════════

def _docx_images(document, api_key, page_est):
    elems = []
    if not api_key: return elems
    try:
        for rel_id, rel in document.part.rels.items():
            if "image" not in rel.reltype: continue
            try:
                ip = rel.target_part; ib = ip.blob; ct = ip.content_type or "image/png"
                if len(ib) < 5000: continue
                cap = caption_image(ib, api_key, page_est, mime=ct)
                if cap:
                    b64 = base64.b64encode(ib).decode("utf-8")
                    elems.append(PageElement(cap, page_est, "image", image_b64=f"data:{ct};base64,{b64}"))
            except: continue
    except: pass
    return elems

def ingest_docx(fp, api_key="", extract_images=True):
    d = docx.Document(fp); elems, sec, lc, ic = [], "", 0, 0
    meta = {"title":d.core_properties.title or "","total_pages":0,
            "file_name":os.path.basename(fp),"file_type":"docx"}
    for p in d.paragraphs:
        txt = p.text.strip()
        if not txt: continue
        lc += max(1, len(txt)//80); pn = lc//45+1
        sty = p.style.name.lower() if p.style else ""
        if "heading" in sty: sec=txt; elems.append(PageElement(txt,pn,"heading",sec))
        else: elems.append(PageElement(txt,pn,"paragraph",sec))
    for tbl in d.tables:
        rows = [[c.text.strip() for c in r.cells] for r in tbl.rows]
        if len(rows)>1:
            hdr="| "+" | ".join(rows[0])+" |"; sep="| "+" | ".join(["---"]*len(rows[0]))+" |"
            body="\n".join("| "+" | ".join(r)+" |" for r in rows[1:])
            elems.append(PageElement(f"{hdr}\n{sep}\n{body}",lc//45+1,"table",sec))
    if extract_images and api_key:
        for ie in _docx_images(d, api_key, lc//45+1):
            ie.section_title=sec; elems.append(ie); ic+=1
    meta["total_pages"]=lc//45+1; meta["images_extracted"]=ic; return elems, meta


# ══════════════════════════════════════════════════════════════
# TXT / Markdown Ingestion
# ══════════════════════════════════════════════════════════════

MD_H=re.compile(r'^(#{1,6})\s+(.+)$'); CAPS_H=re.compile(r'^[A-Z][A-Z\s\d:]{5,80}$')
NUM_H=re.compile(r'^\d+(\.\d+)*\.?\s+[A-Z].{3,80}$')
UL1=re.compile(r'^={3,}\s*$'); UL2=re.compile(r'^-{3,}\s*$')

def ingest_text(fp, is_md=False):
    with open(fp,"r",encoding="utf-8",errors="replace") as f: raw=f.read()
    lines=raw.split("\n"); elems,sec,lc,buf=[],"",[],[]
    meta={"title":"","total_pages":0,"file_name":os.path.basename(fp),
          "file_type":"md" if is_md else "txt","images_extracted":0}
    lc_n=0
    def flush():
        nonlocal buf
        if buf:
            t="\n".join(buf).strip()
            if t: elems.append(PageElement(t,lc_n//50+1,"paragraph",sec))
            buf=[]
    for i,line in enumerate(lines):
        s=line.strip(); lc_n+=1; heading=False
        if not s: flush(); continue
        if is_md:
            m=MD_H.match(s)
            if m: flush();sec=m.group(2).strip();elems.append(PageElement(sec,lc_n//50+1,"heading",sec));heading=True
        if not heading and i+1<len(lines):
            nx=lines[i+1].strip()
            if UL1.match(nx) or UL2.match(nx): flush();sec=s;elems.append(PageElement(sec,lc_n//50+1,"heading",sec));heading=True
        if not heading and not is_md:
            if CAPS_H.match(s) or NUM_H.match(s): flush();sec=s;elems.append(PageElement(sec,lc_n//50+1,"heading",sec));heading=True
        if not heading: buf.append(line)
    flush(); meta["total_pages"]=lc_n//50+1; return elems, meta


def ingest_document(filepath, api_key="", extract_images=True):
    ext = Path(filepath).suffix.lower()
    if ext==".pdf": return ingest_pdf(filepath, api_key, extract_images)
    elif ext==".docx": return ingest_docx(filepath, api_key, extract_images)
    elif ext==".md": return ingest_text(filepath, is_md=True)
    elif ext==".txt": return ingest_text(filepath)
    else: raise ValueError(f"Unsupported: {ext}")


# ══════════════════════════════════════════════════════════════
# Chunking
# ══════════════════════════════════════════════════════════════

tokenizer = tiktoken.encoding_for_model("gpt-4o-mini")
def count_tokens(text): return len(tokenizer.encode(text))

def _make_chunk(texts, pages, sec, idx, has_img=False):
    combined = "\n\n".join(texts)
    cid = hashlib.md5(f"{idx}:{combined[:50]}".encode()).hexdigest()[:10]
    return Chunk(cid, combined, list(pages), sec, count_tokens(combined), idx, has_img)

def chunk_elements(elements, cs=512, ov=64):
    chunks, ci = [], 0
    secs, ct, ce = [], "", []
    for e in elements:
        if e.element_type=="heading":
            if ce: secs.append((ct, ce))
            ct=e.text.strip(); ce=[e]
        else: ce.append(e)
    if ce: secs.append((ct, ce))
    for st, se in secs:
        bt, bp, bk, bi = [], [], 0, False
        for e in se:
            et=count_tokens(e.text); ii=e.element_type=="image"
            if et>cs:
                if bt: chunks.append(_make_chunk(bt,bp,st,ci,bi)); ci+=1; bt,bp,bk,bi=[],[],0,False
                toks=tokenizer.encode(e.text)
                for i in range(0,len(toks),cs):
                    chunks.append(_make_chunk([tokenizer.decode(toks[i:i+cs])],[e.page_num],st,ci,ii)); ci+=1
                continue
            if bk+et>cs and bt:
                chunks.append(_make_chunk(bt,bp,st,ci,bi)); ci+=1
                kt,kp,kk=[],[],0
                for t,p in zip(reversed(bt),reversed(bp)):
                    tt=count_tokens(t)
                    if kk+tt>ov: break
                    kt.insert(0,t);kp.insert(0,p);kk+=tt
                bt,bp,bk,bi=kt,kp,kk,False
            bt.append(e.text);bp.append(e.page_num);bk+=et
            if ii: bi=True
        if bt: chunks.append(_make_chunk(bt,bp,st,ci,bi)); ci+=1
    return [c for c in chunks if c.token_count>=30]


# ══════════════════════════════════════════════════════════════
# Hybrid Index
# ══════════════════════════════════════════════════════════════

class HybridIndex:
    def __init__(self, em):
        self.em=em; self.chunks=[]; self.fi=None; self.bm=None
    def build(self, chunks):
        self.chunks=chunks; texts=[c.text for c in chunks]
        embs=self.em.encode(texts, show_progress_bar=False, batch_size=32, normalize_embeddings=True)
        self.fi=faiss.IndexFlatIP(embs.shape[1]); self.fi.add(embs.astype(np.float32))
        self.bm=BM25Okapi([self._t(t) for t in texts])
    @staticmethod
    def _t(text): return [t for t in re.sub(r'[^a-z0-9\s]',' ',text.lower()).split() if len(t)>1]
    def search(self, q, top_k=15, rrf_k=60):
        qe=self.em.encode([q],normalize_embeddings=True).astype(np.float32)
        _,ids=self.fi.search(qe,top_k)
        dense=[(int(i),r) for r,i in enumerate(ids[0]) if i>=0]
        sc=self.bm.get_scores(self._t(q))
        sparse=[(int(i),r) for r,i in enumerate(np.argsort(sc)[::-1][:top_k]) if sc[i]>0]
        rrf={}
        for i,r in dense: rrf[i]=rrf.get(i,0)+1.0/(rrf_k+r+1)
        for i,r in sparse: rrf[i]=rrf.get(i,0)+1.0/(rrf_k+r+1)
        return [(self.chunks[i],s) for i,s in sorted(rrf.items(),key=lambda x:x[1],reverse=True)[:top_k]]


# ══════════════════════════════════════════════════════════════
# Cross-Encoder Re-Ranker
# ══════════════════════════════════════════════════════════════

class Reranker:
    """
    Second-stage re-ranker using a cross-encoder model.

    First-stage retrieval (FAISS + BM25) optimizes RECALL — cast a wide net.
    The re-ranker optimizes PRECISION — pick the best passages from that net.

    A cross-encoder jointly encodes (query, passage) pairs and produces
    a fine-grained relevance score. This is much more accurate than
    bi-encoder similarity (FAISS) because it attends to token-level
    interactions between query and passage.

    Pipeline: FAISS + BM25 → RRF top-20 → re-rank → top-6 → LLM
    """

    def __init__(self, model_name="cross-encoder/ms-marco-MiniLM-L-6-v2"):
        print(f"Loading re-ranker: {model_name}...")
        self.model = CrossEncoder(model_name, max_length=512)
        print("✅ Re-ranker loaded.")

    def rerank(self, query: str, candidates: List[Tuple[Chunk, float]],
               top_k: int = 6) -> List[Tuple[Chunk, float]]:
        """
        Re-rank candidate chunks using the cross-encoder.

        Args:
            query: User question
            candidates: [(Chunk, rrf_score), ...] from first-stage retrieval
            top_k: Number of passages to return after re-ranking

        Returns:
            Re-ranked [(Chunk, cross_encoder_score), ...] — top_k best
        """
        if not candidates:
            return []

        # Build (query, passage) pairs for the cross-encoder
        pairs = [(query, chunk.text) for chunk, _ in candidates]

        # Score all pairs — cross-encoder sees both query and passage together
        scores = self.model.predict(pairs, show_progress_bar=False)

        # Sort by cross-encoder score (higher = more relevant)
        scored = sorted(zip(candidates, scores), key=lambda x: x[1], reverse=True)

        return [(chunk, float(ce_score)) for (chunk, _), ce_score in scored[:top_k]]


# ══════════════════════════════════════════════════════════════
# Answer Generation
# ══════════════════════════════════════════════════════════════

SYSTEM_PROMPT = """You are a precise document analysis assistant. Answer questions based
ONLY on the provided source passages.

RULES:
1. Use ONLY information from the provided sources. No external knowledge.
2. Cite sources using [Source N] notation.
3. If multiple sources support a claim, cite all: [Source 1, Source 3].
4. If the sources don't contain enough information, say:
   "The provided document sections do not contain sufficient information to answer this question."
5. Do not speculate beyond what the sources explicitly state.
6. For numbers, quote exact figures from the sources.
7. Be concise but complete.
8. Some sources contain descriptions of images, charts, or graphs.
   Treat these as factual content and cite them normally."""

def generate_answer(query, passages, api_key, model="gpt-4o-mini"):
    client = OpenAI(api_key=api_key)
    ctx = "\n\n---\n\n".join(
        f"[Source {i+1}] {c.citation()}\n{c.text}" for i,(c,_) in enumerate(passages))
    resp = client.chat.completions.create(
        model=model, temperature=0.0, max_tokens=1024,
        messages=[
            {"role":"system","content":SYSTEM_PROMPT},
            {"role":"user","content":f"SOURCE PASSAGES:\n{ctx}\n\nQUESTION: {query}\n\nProvide a precise, well-cited answer."}])
    return {
        "answer": resp.choices[0].message.content,
        "sources": [{
            "source_num":i+1, "citation":c.citation(),
            "text":c.text[:300]+("..." if len(c.text)>300 else ""),
            "score":round(s,4), "has_image":c.has_image,
        } for i,(c,s) in enumerate(passages)],
        "usage":{"prompt_tokens":resp.usage.prompt_tokens,
                 "completion_tokens":resp.usage.completion_tokens},
    }


# ══════════════════════════════════════════════════════════════
# Flask App
# ══════════════════════════════════════════════════════════════

app = Flask(__name__)
app.secret_key = os.urandom(24)
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024

# Pipeline config
PIPELINE_CONFIG = {
    "top_k_retrieval": 20,      # candidates from first-stage retrieval
    "top_k_rerank": 6,          # passages after re-ranking → sent to LLM
    "use_reranker": True,       # toggle re-ranker on/off
}

state = {"embed_model":None, "reranker":None, "index":None,
         "metadata":None, "num_chunks":0, "ready":False}
UPLOAD_DIR = os.path.join(tempfile.gettempdir(), "docqa_uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

def get_embed_model():
    if state["embed_model"] is None:
        print("Loading embedding model...")
        state["embed_model"] = SentenceTransformer("BAAI/bge-base-en-v1.5")
        print("✅ Embedding model loaded.")
    return state["embed_model"]

def get_reranker():
    if state["reranker"] is None and PIPELINE_CONFIG["use_reranker"]:
        state["reranker"] = Reranker()
    return state["reranker"]

@app.route("/")
def home(): return render_template("index.html")

@app.route("/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        return jsonify({"error":"No file uploaded"}), 400
    file = request.files["file"]
    ft = get_file_type(file.filename)
    if not ft:
        return jsonify({"error":f"Unsupported. Accepted: {', '.join(SUPPORTED_EXTENSIONS)}"}), 400
    filepath = os.path.join(UPLOAD_DIR, file.filename)
    file.save(filepath)
    api_key = os.environ.get("OPENAI_API_KEY","")
    try:
        t0 = time.time()
        elements, metadata = ingest_document(filepath, api_key=api_key, extract_images=True)
        chunks = chunk_elements(elements, cs=512, ov=64)
        index = HybridIndex(get_embed_model())
        index.build(chunks)
        # Pre-load reranker so first query doesn't pay the load cost
        get_reranker()
        state.update({"index":index,"metadata":metadata,"num_chunks":len(chunks),"ready":True})
        types = {}
        for e in elements: types[e.element_type]=types.get(e.element_type,0)+1
        return jsonify({
            "success":True, "filename":file.filename, "file_type":ft,
            "pages":metadata["total_pages"], "elements":len(elements),
            "chunks":len(chunks), "images":metadata.get("images_extracted",0),
            "element_types":types, "time":round(time.time()-t0,1),
            "reranker":"on" if PIPELINE_CONFIG["use_reranker"] else "off",
        })
    except Exception as e:
        return jsonify({"error":str(e)}), 500

@app.route("/ask", methods=["POST"])
def ask_endpoint():
    if not state["ready"]:
        return jsonify({"error":"No document loaded."}), 400
    query = request.get_json().get("question","").strip()
    if not query:
        return jsonify({"error":"No question provided"}), 400
    api_key = os.environ.get("OPENAI_API_KEY","")
    if not api_key:
        return jsonify({"error":"OPENAI_API_KEY not set"}), 500
    try:
        t0 = time.time()

        # Stage 1: Hybrid retrieval (wide net — recall)
        candidates = state["index"].search(query,
                                           top_k=PIPELINE_CONFIG["top_k_retrieval"])

        # Stage 2: Cross-encoder re-ranking (narrow — precision)
        reranker = get_reranker()
        if reranker:
            passages = reranker.rerank(query, candidates,
                                       top_k=PIPELINE_CONFIG["top_k_rerank"])
        else:
            passages = candidates[:PIPELINE_CONFIG["top_k_rerank"]]

        # Stage 3: Generate answer
        result = generate_answer(query, passages, api_key)
        result.update({"latency":round(time.time()-t0,2), "question":query,
                       "reranked": reranker is not None})
        return jsonify(result)
    except Exception as e:
        return jsonify({"error":str(e)}), 500

@app.route("/status")
def status():
    if state["ready"]:
        return jsonify({"ready":True, "filename":state["metadata"]["file_name"],
                        "pages":state["metadata"]["total_pages"],
                        "chunks":state["num_chunks"],
                        "images":state["metadata"].get("images_extracted",0),
                        "reranker":"on" if PIPELINE_CONFIG["use_reranker"] else "off"})
    return jsonify({"ready":False})

if __name__ == "__main__":
    print("\n" + "="*50)
    print("  Document Q&A — Multimodal + Re-Ranker")
    print("  Formats: PDF, DOCX, TXT, Markdown")
    print("  Images: PDF + DOCX")
    print(f"  Re-ranker: {'ON' if PIPELINE_CONFIG['use_reranker'] else 'OFF'}")
    print("="*50)
    print(f"  http://localhost:5000")
    print(f"  API key: {'✅' if os.environ.get('OPENAI_API_KEY') else '❌'}")
    print("="*50 + "\n")
    # Pre-load models on startup
    get_embed_model()
    get_reranker()
    app.run(debug=True, port=5000)
