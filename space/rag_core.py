"""
rag_core.py — Core multimodal RAG pipeline (single source of truth).

Imported by:
  • app.py            — Flask web app (local, full-featured)
  • eval_compare.py   — FinanceBench A/B benchmark
  • space/app.py      — Gradio Hugging Face Space (public demo)

Pipeline: ingest → chunk → hybrid retrieve (BGE + BM25 + CLIP) → re-rank → generate.

Two multimodal upgrades over a plain caption-then-embed RAG, both toggleable:
  #1 use_image_in_answer  — feed the *actual image pixels* to the VLM at answer
                            time, not just the caption text. The model can
                            re-read a chart for the specific question asked.
  #2 use_clip_retrieval   — embed image pixels with CLIP and fold a visual
                            retrieval leg into Reciprocal Rank Fusion, so a
                            figure is findable by its content even when the
                            caption missed the detail the user asked about.
"""

import os, re, io, base64, hashlib, json
from dataclasses import dataclass, field
from typing import List, Optional, Tuple

import numpy as np

# ── Optional heavy deps (imported lazily where possible) ──────────────────────
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
try:
    import docx  # python-docx
except ImportError:
    docx = None


# ══════════════════════════════════════════════════════════════
# Configuration
# ══════════════════════════════════════════════════════════════

DEFAULT_CONFIG = {
    "chunk_size": 512,
    "chunk_overlap": 64,
    "embedding_model": "BAAI/bge-base-en-v1.5",
    "reranker_model": "cross-encoder/ms-marco-MiniLM-L-6-v2",
    "clip_model": "clip-ViT-B-32",          # runs on CPU (~600 MB)
    "use_reranker": True,
    "top_k_retrieval": 20,
    "top_k_rerank": 6,
    "rrf_k": 60,
    "llm_model": "gpt-4o-mini",
    "caption_model": "gpt-4o-mini",
    "temperature": 0.0,
    "max_tokens": 1024,
    "extract_images": True,
    # ── multimodal upgrades ──
    "use_image_in_answer": True,            # #1
    "use_clip_retrieval": True,             # #2
    "max_answer_images": 3,                 # cost cap: images sent to the VLM
    "image_detail": "auto",                 # "low" is cheaper; "high" reads charts better
}


def make_config(**overrides):
    cfg = dict(DEFAULT_CONFIG)
    cfg.update(overrides)
    return cfg


# ══════════════════════════════════════════════════════════════
# Data structures
# ══════════════════════════════════════════════════════════════

@dataclass
class PageElement:
    text: str
    page_num: int
    element_type: str
    section_title: str = ""
    image_b64: str = ""          # data URI: "data:<mime>;base64,<...>"


@dataclass
class Chunk:
    chunk_id: str
    text: str
    page_numbers: List[int]
    section_title: str
    token_count: int
    chunk_index: int
    has_image: bool = False
    image_b64s: List[str] = field(default_factory=list)   # data URIs of images in this chunk (#1, #2)

    def citation(self):
        pages = sorted(set(self.page_numbers))
        p = f"p. {pages[0]}" if len(pages) == 1 else f"pp. {pages[0]}-{pages[-1]}"
        sec = f' — "{self.section_title}"' if self.section_title else ""
        img = " [image]" if self.has_image else ""
        return f"[{p}{sec}{img}]"


# ══════════════════════════════════════════════════════════════
# Lazy, cached model loaders (shared across app / eval / space)
# ══════════════════════════════════════════════════════════════

_MODELS = {}


def get_embed_model(name=DEFAULT_CONFIG["embedding_model"]):
    key = ("embed", name)
    if key not in _MODELS:
        from sentence_transformers import SentenceTransformer
        _MODELS[key] = SentenceTransformer(name)
    return _MODELS[key]


def get_clip_model(name=DEFAULT_CONFIG["clip_model"]):
    key = ("clip", name)
    if key not in _MODELS:
        from sentence_transformers import SentenceTransformer
        _MODELS[key] = SentenceTransformer(name)
    return _MODELS[key]


def get_reranker(name=DEFAULT_CONFIG["reranker_model"]):
    key = ("rerank", name)
    if key not in _MODELS:
        _MODELS[key] = Reranker(name)
    return _MODELS[key]


# ── tokenizer (tiktoken) ──────────────────────────────────────────────────────
_tokenizer = None


def _tok():
    global _tokenizer
    if _tokenizer is None:
        import tiktoken
        _tokenizer = tiktoken.encoding_for_model("gpt-4o-mini")
    return _tokenizer


def count_tokens(text: str) -> int:
    return len(_tok().encode(text))


# ── base64 / image helpers ────────────────────────────────────────────────────
def data_uri(image_bytes: bytes, mime: str) -> str:
    return f"data:{mime};base64,{base64.b64encode(image_bytes).decode('utf-8')}"


def decode_data_uri(uri: str) -> bytes:
    return base64.b64decode(uri.split(",", 1)[1])


def _pil_from_uri(uri: str):
    from PIL import Image
    return Image.open(io.BytesIO(decode_data_uri(uri))).convert("RGB")


# ══════════════════════════════════════════════════════════════
# Caption cache — bound API cost by never re-captioning the same image
# ══════════════════════════════════════════════════════════════

class CaptionCache:
    """Disk-backed cache of image captions keyed by (image bytes, prompt, model)."""

    def __init__(self, path: Optional[str] = None):
        self.path = path
        self.store = {}
        if path and os.path.isfile(path):
            try:
                self.store = json.load(open(path))
            except Exception:
                self.store = {}

    @staticmethod
    def _key(image_bytes, prompt, model):
        h = hashlib.sha256(image_bytes).hexdigest()
        return f"{model}:{hashlib.md5(prompt.encode()).hexdigest()[:8]}:{h}"

    def get(self, image_bytes, prompt, model):
        return self.store.get(self._key(image_bytes, prompt, model))

    def put(self, image_bytes, prompt, model, caption):
        self.store[self._key(image_bytes, prompt, model)] = caption
        if self.path:
            try:
                json.dump(self.store, open(self.path, "w"))
            except Exception:
                pass


# ══════════════════════════════════════════════════════════════
# Image captioning (the VLM step — gpt-4o-mini is already vision-capable)
# ══════════════════════════════════════════════════════════════

IMAGE_CAPTION_PROMPT = (
    "Describe this image from a document in detail. "
    "If it is a chart/graph/table, describe data, trends, labels, axes, key values. "
    "Be specific about numbers, percentages, or text visible.")

MIME_MAP = {"png": "image/png", "jpeg": "image/jpeg", "jpg": "image/jpeg",
            "gif": "image/gif", "bmp": "image/bmp", "tiff": "image/tiff", "webp": "image/webp"}


def caption_image(image_bytes, api_key, page_num, model="gpt-4o-mini",
                  mime="image/png", cache: Optional[CaptionCache] = None):
    if not api_key or len(image_bytes) < 5000:
        return None
    if cache is not None:
        hit = cache.get(image_bytes, IMAGE_CAPTION_PROMPT, model)
        if hit is not None:
            return f"[Image on page {page_num}]: {hit}"
    from openai import OpenAI
    b64 = base64.b64encode(image_bytes).decode("utf-8")
    try:
        resp = OpenAI(api_key=api_key).chat.completions.create(
            model=model, temperature=0.0, max_tokens=300,
            messages=[{"role": "user", "content": [
                {"type": "text", "text": IMAGE_CAPTION_PROMPT},
                {"type": "image_url",
                 "image_url": {"url": f"data:{mime};base64,{b64}", "detail": "high"}}]}])
        text = resp.choices[0].message.content
        if cache is not None:
            cache.put(image_bytes, IMAGE_CAPTION_PROMPT, model, text)
        return f"[Image on page {page_num}]: {text}"
    except Exception as e:
        print(f"  ⚠️ Caption failed p.{page_num}: {e}")
        return None


# ══════════════════════════════════════════════════════════════
# PDF ingestion
# ══════════════════════════════════════════════════════════════

def _pdf_images(doc, page, pn, api_key, model, cache, budget=None, min_w=100, min_h=100):
    """budget: a mutable [remaining] counter capping total captions, or None for no cap."""
    elems = []
    for info in page.get_images(full=True):
        if budget is not None and budget[0] <= 0:
            break
        try:
            bi = doc.extract_image(info[0])
            if not bi or bi.get("width", 0) < min_w or bi.get("height", 0) < min_h:
                continue
            ext = bi.get("ext", "png"); mime = MIME_MAP.get(ext, "image/png")
            cap = caption_image(bi["image"], api_key, pn, model=model, mime=mime, cache=cache)
            if cap:
                elems.append(PageElement(cap, pn, "image", image_b64=data_uri(bi["image"], mime)))
                if budget is not None:
                    budget[0] -= 1
        except Exception:
            continue
    return elems


def _median_fs(page):
    sizes = []
    for b in page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]:
        if b.get("type", 1) != 0:
            continue
        for l in b.get("lines", []):
            for s in l.get("spans", []):
                t = s.get("text", "").strip()
                if len(t) > 2:
                    sizes.extend([s["size"]] * len(t))
    return float(np.median(sizes)) if sizes else 12.0


def _is_heading_pdf(block, med):
    lines = block.get("lines", [])
    if not lines:
        return False
    fs, bold = [], False
    for l in lines:
        for s in l.get("spans", []):
            fs.append(s["size"])
            if "bold" in s.get("font", "").lower():
                bold = True
    if not fs:
        return False
    r = max(fs) / med if med > 0 else 1.0
    txt = " ".join(s["text"] for l in lines for s in l.get("spans", [])).strip()
    if len(txt) > 200 or len(txt) < 2:
        return False
    return r >= 1.5 or (r >= 1.3 and bold) or r >= 1.2 or (bold and len(txt) < 80)


def ingest_pdf(fp, api_key="", extract_images=True, caption_model="gpt-4o-mini",
               cache=None, max_caption_images=None):
    doc = fitz.open(fp); elems, sec, ic = [], "", 0
    budget = [max_caption_images] if max_caption_images is not None else None
    meta = {"title": doc.metadata.get("title", ""), "total_pages": len(doc),
            "file_name": os.path.basename(fp), "file_type": "pdf"}
    for pi in range(len(doc)):
        page = doc[pi]; pn = pi + 1; med = _median_fs(page)
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        try:
            for t in page.find_tables():
                df = t.to_pandas()
                if not df.empty:
                    elems.append(PageElement(df.to_markdown(index=False), pn, "table", sec))
        except Exception:
            pass
        for b in blocks:
            if b.get("type", 1) != 0:
                continue
            txt = ""
            for l in b.get("lines", []):
                txt += "".join(s["text"] for s in l.get("spans", [])) + "\n"
            txt = txt.strip()
            if not txt or len(txt) < 3:
                continue
            if _is_heading_pdf(b, med):
                sec = txt.strip(); elems.append(PageElement(txt, pn, "heading", sec))
            else:
                elems.append(PageElement(txt, pn, "paragraph", sec))
        if extract_images and api_key and (budget is None or budget[0] > 0):
            for ie in _pdf_images(doc, page, pn, api_key, caption_model, cache, budget):
                ie.section_title = sec; elems.append(ie); ic += 1
    doc.close(); meta["images_extracted"] = ic
    return elems, meta


# ══════════════════════════════════════════════════════════════
# DOCX ingestion
# ══════════════════════════════════════════════════════════════

def _docx_images(document, api_key, page_est, caption_model, cache, budget=None):
    elems = []
    if not api_key:
        return elems
    try:
        for rel_id, rel in document.part.rels.items():
            if "image" not in rel.reltype:
                continue
            if budget is not None and budget[0] <= 0:
                break
            try:
                ip = rel.target_part; ib = ip.blob; ct = ip.content_type or "image/png"
                if len(ib) < 5000:
                    continue
                cap = caption_image(ib, api_key, page_est, model=caption_model, mime=ct, cache=cache)
                if cap:
                    elems.append(PageElement(cap, page_est, "image", image_b64=data_uri(ib, ct)))
                    if budget is not None:
                        budget[0] -= 1
            except Exception:
                continue
    except Exception:
        pass
    return elems


def ingest_docx(fp, api_key="", extract_images=True, caption_model="gpt-4o-mini",
                cache=None, max_caption_images=None):
    d = docx.Document(fp); elems, sec, lc, ic = [], "", 0, 0
    budget = [max_caption_images] if max_caption_images is not None else None
    meta = {"title": d.core_properties.title or "", "total_pages": 0,
            "file_name": os.path.basename(fp), "file_type": "docx"}
    for p in d.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        lc += max(1, len(txt) // 80); pn = lc // 45 + 1
        sty = p.style.name.lower() if p.style else ""
        if "heading" in sty:
            sec = txt; elems.append(PageElement(txt, pn, "heading", sec))
        else:
            elems.append(PageElement(txt, pn, "paragraph", sec))
    for tbl in d.tables:
        rows = [[c.text.strip() for c in r.cells] for r in tbl.rows]
        if len(rows) > 1:
            hdr = "| " + " | ".join(rows[0]) + " |"
            sep = "| " + " | ".join(["---"] * len(rows[0])) + " |"
            body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
            elems.append(PageElement(f"{hdr}\n{sep}\n{body}", lc // 45 + 1, "table", sec))
    if extract_images and api_key:
        for ie in _docx_images(d, api_key, lc // 45 + 1, caption_model, cache, budget):
            ie.section_title = sec; elems.append(ie); ic += 1
    meta["total_pages"] = lc // 45 + 1; meta["images_extracted"] = ic
    return elems, meta


# ══════════════════════════════════════════════════════════════
# TXT / Markdown ingestion
# ══════════════════════════════════════════════════════════════

MD_H = re.compile(r'^(#{1,6})\s+(.+)$'); CAPS_H = re.compile(r'^[A-Z][A-Z\s\d:]{5,80}$')
NUM_H = re.compile(r'^\d+(\.\d+)*\.?\s+[A-Z].{3,80}$')
UL1 = re.compile(r'^={3,}\s*$'); UL2 = re.compile(r'^-{3,}\s*$')


def ingest_text(fp, is_md=False):
    with open(fp, "r", encoding="utf-8", errors="replace") as f:
        raw = f.read()
    lines = raw.split("\n"); elems, sec, buf = [], "", []
    meta = {"title": "", "total_pages": 0, "file_name": os.path.basename(fp),
            "file_type": "md" if is_md else "txt", "images_extracted": 0}
    lc_n = 0

    def flush():
        if buf:
            t = "\n".join(buf).strip()
            if t:
                elems.append(PageElement(t, lc_n // 50 + 1, "paragraph", sec))
            buf.clear()

    for i, line in enumerate(lines):
        s = line.strip(); lc_n += 1; heading = False
        if not s:
            flush(); continue
        if is_md:
            m = MD_H.match(s)
            if m:
                flush(); sec = m.group(2).strip()
                elems.append(PageElement(sec, lc_n // 50 + 1, "heading", sec)); heading = True
        if not heading and i + 1 < len(lines):
            nx = lines[i + 1].strip()
            if UL1.match(nx) or UL2.match(nx):
                flush(); sec = s
                elems.append(PageElement(sec, lc_n // 50 + 1, "heading", sec)); heading = True
        if not heading and not is_md:
            if CAPS_H.match(s) or NUM_H.match(s):
                flush(); sec = s
                elems.append(PageElement(sec, lc_n // 50 + 1, "heading", sec)); heading = True
        if not heading:
            buf.append(line)
    flush(); meta["total_pages"] = lc_n // 50 + 1
    return elems, meta


def ingest_document(filepath, api_key="", extract_images=True,
                    caption_model="gpt-4o-mini", cache=None, max_caption_images=None):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return ingest_pdf(filepath, api_key, extract_images, caption_model, cache, max_caption_images)
    elif ext == ".docx":
        return ingest_docx(filepath, api_key, extract_images, caption_model, cache, max_caption_images)
    elif ext == ".md":
        return ingest_text(filepath, is_md=True)
    elif ext == ".txt":
        return ingest_text(filepath)
    else:
        raise ValueError(f"Unsupported: {ext}")


# ══════════════════════════════════════════════════════════════
# Chunking
# ══════════════════════════════════════════════════════════════

def _make_chunk(texts, pages, sec, idx, has_img=False, image_uris=None):
    combined = "\n\n".join(texts)
    cid = hashlib.md5(f"{idx}:{combined[:50]}".encode()).hexdigest()[:10]
    return Chunk(cid, combined, list(pages), sec, count_tokens(combined), idx,
                 has_img, list(image_uris or []))


def chunk_elements(elements, cs=512, ov=64):
    chunks, ci = [], 0
    secs, ct, ce = [], "", []
    for e in elements:
        if e.element_type == "heading":
            if ce:
                secs.append((ct, ce))
            ct = e.text.strip(); ce = [e]
        else:
            ce.append(e)
    if ce:
        secs.append((ct, ce))
    for st, se in secs:
        bt, bp, bk, bi, bimg = [], [], 0, False, []
        for e in se:
            et = count_tokens(e.text); is_img = e.element_type == "image"
            if et > cs:
                if bt:
                    chunks.append(_make_chunk(bt, bp, st, ci, bi, bimg)); ci += 1
                    bt, bp, bk, bi, bimg = [], [], 0, False, []
                toks = _tok().encode(e.text)
                for i in range(0, len(toks), cs):
                    chunks.append(_make_chunk([_tok().decode(toks[i:i + cs])], [e.page_num],
                                              st, ci, is_img, [e.image_b64] if is_img and e.image_b64 else []))
                    ci += 1
                continue
            if bk + et > cs and bt:
                chunks.append(_make_chunk(bt, bp, st, ci, bi, bimg)); ci += 1
                kt, kp, kk = [], [], 0
                for t, p in zip(reversed(bt), reversed(bp)):
                    tt = count_tokens(t)
                    if kk + tt > ov:
                        break
                    kt.insert(0, t); kp.insert(0, p); kk += tt
                bt, bp, bk, bi, bimg = kt, kp, kk, False, []
            bt.append(e.text); bp.append(e.page_num); bk += et
            if is_img:
                bi = True
                if e.image_b64:
                    bimg.append(e.image_b64)
        if bt:
            chunks.append(_make_chunk(bt, bp, st, ci, bi, bimg)); ci += 1
    return [c for c in chunks if c.token_count >= 30]


# ══════════════════════════════════════════════════════════════
# Hybrid index — BGE dense + BM25 sparse + CLIP visual (#2), fused via RRF
# ══════════════════════════════════════════════════════════════

# Tokenizer that PRESERVES financial tokens ($, %, decimals, parenthesised
# negatives) so BM25 can match numeric queries — critical on FinanceBench.
_NUM = re.compile(r'\$?\(?-?\d[\d,]*\.?\d*%?\)?')
_WORD = re.compile(r'[a-z0-9]+')


def bm25_tokenize(text: str) -> List[str]:
    text_l = text.lower()
    nums = [m.group().strip("()") for m in _NUM.finditer(text_l)]
    words = [w for w in _WORD.findall(re.sub(r'[^a-z0-9\s]', ' ', text_l)) if len(w) > 1]
    return words + nums


class HybridIndex:
    """
    Three retrieval legs fused with Reciprocal Rank Fusion:
      • dense   — BGE bi-encoder over chunk text (semantic recall)
      • sparse  — BM25 over numeric-aware tokens (exact terms / figures)
      • visual  — CLIP query-text → image-pixel similarity (#2), over the
                  subset of chunks that carry images. Lets a figure surface
                  even when its caption missed the asked-for detail.
    """

    def __init__(self, embed_model, clip_model=None):
        self.em = embed_model
        self.clip = clip_model        # None disables the visual leg
        self.chunks = []
        self.fi = None
        self.bm = None
        self.img_fi = None
        self.img_chunk_ids = []       # row in img_fi → index into self.chunks

    def build(self, chunks):
        import faiss
        from rank_bm25 import BM25Okapi
        self.chunks = chunks
        texts = [c.text for c in chunks]
        embs = self.em.encode(texts, show_progress_bar=False, batch_size=32,
                              normalize_embeddings=True)
        self.fi = faiss.IndexFlatIP(embs.shape[1]); self.fi.add(embs.astype(np.float32))
        self.bm = BM25Okapi([bm25_tokenize(t) for t in texts])
        if self.clip is not None:
            self._build_image_index(faiss)

    def _build_image_index(self, faiss):
        vecs, ids = [], []
        for ci, c in enumerate(self.chunks):
            if not c.image_b64s:
                continue
            imgs = []
            for uri in c.image_b64s:
                try:
                    imgs.append(_pil_from_uri(uri))
                except Exception:
                    continue
            if not imgs:
                continue
            v = self.clip.encode(imgs, show_progress_bar=False, normalize_embeddings=True)
            vecs.append(np.asarray(v, dtype=np.float32).mean(axis=0))  # one vector per chunk
            ids.append(ci)
        if vecs:
            mat = np.vstack(vecs).astype(np.float32)
            faiss.normalize_L2(mat)
            self.img_fi = faiss.IndexFlatIP(mat.shape[1]); self.img_fi.add(mat)
            self.img_chunk_ids = ids

    def search(self, q, top_k=15, rrf_k=60):
        qe = self.em.encode([q], normalize_embeddings=True).astype(np.float32)
        _, ids = self.fi.search(qe, top_k)
        dense = [(int(i), r) for r, i in enumerate(ids[0]) if i >= 0]

        sc = self.bm.get_scores(bm25_tokenize(q))
        sparse = [(int(i), r) for r, i in enumerate(np.argsort(sc)[::-1][:top_k]) if sc[i] > 0]

        visual = []
        if self.clip is not None and self.img_fi is not None:
            qv = self.clip.encode([q], normalize_embeddings=True).astype(np.float32)
            k = min(top_k, self.img_fi.ntotal)
            _, vids = self.img_fi.search(qv, k)
            visual = [(self.img_chunk_ids[int(i)], r) for r, i in enumerate(vids[0]) if i >= 0]

        rrf = {}
        for legs in (dense, sparse, visual):
            for i, r in legs:
                rrf[i] = rrf.get(i, 0) + 1.0 / (rrf_k + r + 1)
        ranked = sorted(rrf.items(), key=lambda x: x[1], reverse=True)[:top_k]
        return [(self.chunks[i], s) for i, s in ranked]


# ══════════════════════════════════════════════════════════════
# Cross-encoder re-ranker
# ══════════════════════════════════════════════════════════════

class Reranker:
    """Second-stage precision re-ranker. Jointly encodes (query, passage) pairs."""

    def __init__(self, model_name="cross-encoder/ms-marco-MiniLM-L-6-v2"):
        from sentence_transformers import CrossEncoder
        self.model = CrossEncoder(model_name, max_length=512)

    def rerank(self, query: str, candidates: List[Tuple[Chunk, float]],
               top_k: int = 6) -> List[Tuple[Chunk, float]]:
        if not candidates:
            return []
        pairs = [(query, chunk.text) for chunk, _ in candidates]
        scores = self.model.predict(pairs, show_progress_bar=False)
        scored = sorted(zip(candidates, scores), key=lambda x: x[1], reverse=True)
        return [(chunk, float(ce)) for (chunk, _), ce in scored[:top_k]]


# ══════════════════════════════════════════════════════════════
# Answer generation — text context + (optionally) the real image pixels (#1)
# ══════════════════════════════════════════════════════════════

SYSTEM_PROMPT = """You are a precise document analysis assistant. Answer questions based
ONLY on the provided source passages.

RULES:
1. Use ONLY information from the provided sources. No external knowledge.
2. Cite sources using [Source N] notation.
3. If multiple sources support a claim, cite all: [Source 1, Source 3].
4. If the sources don't contain enough information, say:
   "The provided document sections do not contain sufficient information to answer this question."
5. Do not speculate beyond what the sources explicitly state.
6. For numbers, quote exact figures from the sources.
7. Be concise but complete.
8. Some sources include an attached image (chart/graph/table). When an image is
   attached for a source, read it directly and treat what you see as factual
   content for that source, cited normally."""


def _build_user_content(query, passages, cfg):
    ctx = "\n\n---\n\n".join(
        f"[Source {i+1}] {c.citation()}\n{c.text}" for i, (c, _) in enumerate(passages))
    text_block = (f"SOURCE PASSAGES:\n{ctx}\n\nQUESTION: {query}\n\n"
                  "Provide a precise, well-cited answer.")

    if not cfg.get("use_image_in_answer"):
        return text_block

    # #1 — attach the actual image pixels for image-bearing sources (cost-capped).
    image_parts, attached = [], []
    budget = cfg.get("max_answer_images", 3)
    for i, (c, _) in enumerate(passages):
        if len(image_parts) >= budget:
            break
        for uri in c.image_b64s:
            if len(image_parts) >= budget:
                break
            image_parts.append({"type": "image_url",
                                "image_url": {"url": uri, "detail": cfg.get("image_detail", "auto")}})
            attached.append(i + 1)
    if not image_parts:
        return text_block

    note = ("\n\nNOTE: The image(s) attached below belong to "
            + ", ".join(f"[Source {n}]" for n in attached)
            + ". Read them directly to answer.")
    return [{"type": "text", "text": text_block + note}] + image_parts


def generate_answer(query, passages, api_key, cfg=None):
    cfg = cfg or DEFAULT_CONFIG
    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    resp = client.chat.completions.create(
        model=cfg["llm_model"], temperature=cfg.get("temperature", 0.0),
        max_tokens=cfg.get("max_tokens", 1024),
        messages=[{"role": "system", "content": SYSTEM_PROMPT},
                  {"role": "user", "content": _build_user_content(query, passages, cfg)}])
    return {
        "query": query,
        "answer": resp.choices[0].message.content,
        "sources": [{"source_num": i + 1, "chunk": c, "citation": c.citation(),
                     "score": round(float(s), 4), "has_image": c.has_image,
                     "text": c.text[:300] + ("..." if len(c.text) > 300 else "")}
                    for i, (c, s) in enumerate(passages)],
        "usage": {"prompt_tokens": resp.usage.prompt_tokens,
                  "completion_tokens": resp.usage.completion_tokens},
    }


def ask(query, index, cfg, reranker=None, api_key=None):
    """Full pipeline: hybrid retrieve → re-rank → generate."""
    api_key = api_key if api_key is not None else os.environ.get("OPENAI_API_KEY", "")
    candidates = index.search(query, top_k=cfg["top_k_retrieval"], rrf_k=cfg["rrf_k"])
    if reranker and cfg.get("use_reranker", True):
        passages = reranker.rerank(query, candidates, top_k=cfg["top_k_rerank"])
    else:
        passages = candidates[:cfg["top_k_rerank"]]
    return generate_answer(query, passages, api_key, cfg)
